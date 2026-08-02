"""Build the formal ECK v0.1 Traditional Chinese engineering specification."""

from __future__ import annotations

import argparse
import json
import re
from collections.abc import Iterable
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = PROJECT_ROOT / "docs"
RELEASE_REPORT = PROJECT_ROOT / "artifacts" / "release-report.json"

PAGE_WIDTH_DXA = 12240
USABLE_WIDTH_DXA = 9360
BODY_FONT = "Noto Sans T Chinese"
LATIN_FONT = "Aptos"
MONO_FONT = "DejaVu Sans Mono"
NAVY = "17324D"
BLUE = "2E74B5"
TEAL = "0F6B6D"
GOLD = "C8922E"
INK = "23313F"
MUTED = "5D6B78"
PALE_BLUE = "E8EEF5"
PALE_TEAL = "E5F2F1"
PALE_GOLD = "F8F0DF"
WHITE = "FFFFFF"
GRID = "B9C5D0"
RELEASE_DATE = "2026-07-29"

VOLUMES = [
    ("I", "Vision, Philosophy & Constitution", "01-vision-philosophy-constitution.md"),
    ("II", "Architecture Specification", "02-architecture-specification.md"),
    ("III", "Digital Life Kernel", "03-digital-life-kernel.md"),
    ("IV", "Memory System", "04-memory-system.md"),
    ("V", "Experience Engine", "05-experience-engine.md"),
    ("VI", "Brain Runtime", "06-brain-runtime.md"),
    ("VII", "Prediction & World Action Model", "07-prediction-world-action-model.md"),
    ("VIII", "Planner, Reflection & Curiosity", "08-planner-reflection-curiosity.md"),
    ("IX", "Contracts & API", "09-contracts-api.md"),
    ("X", "Runtime & Infrastructure", "10-runtime-infrastructure.md"),
    ("XI", "Development Guide", "11-development-guide.md"),
    ("XII", "Testing & Validation", "12-testing-validation.md"),
    ("XIII", "Roadmap & Research", "13-roadmap-research.md"),
]

ADR_FILES = [
    "0001-system-not-model.md",
    "0002-verifier-grounded-learning.md",
    "0003-sqlite-event-chain.md",
    "0004-no-weight-update-v01.md",
    "0005-no-arbitrary-code.md",
    "0006-ollama-default-provider.md",
    "0007-python-first-rust-verifier.md",
]

INLINE_TOKEN = re.compile(
    r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|<https?://[^>]+>|\*[^*]+\*)"
)


def set_cell_shading(cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(
    cell: _Cell,
    *,
    top: int = 80,
    start: int = 120,
    bottom: int = 80,
    end: int = 120,
) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = tc_mar.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_mar.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")


def set_table_borders(table: Table, color: str = GRID, size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def remove_table_borders(table: Table) -> None:
    set_table_borders(table, color=WHITE, size=0)
    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    if borders is not None:
        for edge in borders:
            edge.set(qn("w:val"), "nil")


def set_table_width(table: Table, width_dxa: int = USABLE_WIDTH_DXA) -> None:
    tbl_pr = table._tbl.tblPr
    width = tbl_pr.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_cell_width(cell: _Cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def repeat_table_header(row: object) -> None:
    tr_pr = row._tr.get_or_add_trPr()  # type: ignore[attr-defined]
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row: object) -> None:
    tr_pr = row._tr.get_or_add_trPr()  # type: ignore[attr-defined]
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_run_font(
    run: Run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
    mono: bool = False,
) -> None:
    font_name = MONO_FONT if mono else LATIN_FONT
    east_asia = MONO_FONT if mono else BODY_FONT
    run.font.name = font_name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    fonts = run._element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    lang = run._element.rPr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        run._element.rPr.append(lang)
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:eastAsia"), "zh-TW")


def add_hyperlink(paragraph: Paragraph, text: str, url: str) -> None:
    rel_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run_element = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), LATIN_FONT)
    fonts.set(qn("w:hAnsi"), LATIN_FONT)
    fonts.set(qn("w:eastAsia"), BODY_FONT)
    properties.extend([fonts, color, underline])
    run_element.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def add_inline(paragraph: Paragraph, text: str, *, size: float = 11, color: str = INK) -> None:
    cursor = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=max(8.5, size - 1), color=TEAL, mono=True)
        elif token.startswith("["):
            label, url = token[1:].split("](", 1)
            add_hyperlink(paragraph, label, url[:-1])
        elif token.startswith("<"):
            url = token[1:-1]
            add_hyperlink(paragraph, url, url)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, italic=True, color=color)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=color)


def add_page_field(paragraph: Paragraph) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    properties.extend([color, size])
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def create_ordered_list_numbering(doc: DocumentType, *, start: int = 1) -> int:
    """Create a document-local decimal list whose first item starts at ``start``."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId"), "0"))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(element.get(qn("w:numId"), "0"))
        for element in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "multilevel")
    abstract.append(multi_level)

    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))

        start_element = OxmlElement("w:start")
        start_element.set(qn("w:val"), str(start if level == 0 else 1))
        lvl.append(start_element)

        num_format = OxmlElement("w:numFmt")
        num_format.set(qn("w:val"), "decimal")
        lvl.append(num_format)

        level_text = OxmlElement("w:lvlText")
        level_text.set(qn("w:val"), f"%{level + 1}.")
        lvl.append(level_text)

        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), "space")
        lvl.append(suffix)

        paragraph_properties = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(540 + level * 360))
        tabs.append(tab)
        indentation = OxmlElement("w:ind")
        indentation.set(qn("w:left"), str(540 + level * 360))
        indentation.set(qn("w:hanging"), "270")
        paragraph_properties.extend([tabs, indentation])
        lvl.append(paragraph_properties)
        abstract.append(lvl)

    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    num.append(abstract_reference)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start))
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def create_bullet_list_numbering(doc: DocumentType) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId"), "0"))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(element.get(qn("w:numId"), "0"))
        for element in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "multilevel")
    abstract.append(multi_level)
    markers = ("•", "◦", "▪")
    for level, marker in enumerate(markers):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start_element = OxmlElement("w:start")
        start_element.set(qn("w:val"), "1")
        num_format = OxmlElement("w:numFmt")
        num_format.set(qn("w:val"), "bullet")
        level_text = OxmlElement("w:lvlText")
        level_text.set(qn("w:val"), marker)
        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), "space")
        paragraph_properties = OxmlElement("w:pPr")
        indentation = OxmlElement("w:ind")
        indentation.set(qn("w:left"), str(540 + level * 360))
        indentation.set(qn("w:hanging"), "270")
        paragraph_properties.append(indentation)
        lvl.extend(
            [
                start_element,
                num_format,
                level_text,
                suffix,
                paragraph_properties,
            ]
        )
        abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    num.append(abstract_reference)
    numbering.append(num)
    return num_id


def apply_ordered_list_numbering(paragraph: Paragraph, num_id: int, level: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    existing = paragraph_properties.find(qn("w:numPr"))
    if existing is not None:
        paragraph_properties.remove(existing)
    number_properties = OxmlElement("w:numPr")
    level_element = OxmlElement("w:ilvl")
    level_element.set(qn("w:val"), str(level))
    id_element = OxmlElement("w:numId")
    id_element.set(qn("w:val"), str(num_id))
    number_properties.extend([level_element, id_element])
    paragraph_properties.append(number_properties)


def set_bottom_border(paragraph: Paragraph, color: str = GRID) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_styles(doc: DocumentType) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, "1F4D78", 10, 5),
        "Heading 4": (11, TEAL, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = LATIN_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = LATIN_FONT
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "ECK Code" not in styles:
        code_style = styles.add_style("ECK Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["ECK Code"]
    code_style.font.name = MONO_FONT
    code_style.font.size = Pt(8)
    code_style._element.rPr.rFonts.set(qn("w:ascii"), MONO_FONT)
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), MONO_FONT)
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    code_style.paragraph_format.left_indent = Inches(0.18)
    code_style.paragraph_format.right_indent = Inches(0.12)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(8)
    code_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def configure_section(section: object) -> None:
    section.page_width = Inches(8.5)  # type: ignore[attr-defined]
    section.page_height = Inches(11)  # type: ignore[attr-defined]
    section.top_margin = Inches(1)  # type: ignore[attr-defined]
    section.bottom_margin = Inches(1)  # type: ignore[attr-defined]
    section.left_margin = Inches(1)  # type: ignore[attr-defined]
    section.right_margin = Inches(1)  # type: ignore[attr-defined]
    section.header_distance = Inches(0.492)  # type: ignore[attr-defined]
    section.footer_distance = Inches(0.492)  # type: ignore[attr-defined]


def configure_header_footer(section: object) -> None:
    header = section.header  # type: ignore[attr-defined]
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("ECK")
    set_run_font(run, size=9, bold=True, color=TEAL)
    run = paragraph.add_run("  |  Embodied Cognitive Kernel")
    set_run_font(run, size=8.5, color=MUTED)
    set_bottom_border(paragraph)

    footer = section.footer  # type: ignore[attr-defined]
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    remove_table_borders(table)
    left, right = table.rows[0].cells
    left.width = Inches(5.7)
    right.width = Inches(0.8)
    p_left = left.paragraphs[0]
    p_left.paragraph_format.space_after = Pt(0)
    run = p_left.add_run("Digital Life Kernel v0.1  •  Apache-2.0  •  2026-07-29")
    set_run_font(run, size=8.5, color=MUTED)
    p_right = right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.space_after = Pt(0)
    add_page_field(p_right)

    pg_num_type = section._sectPr.find(qn("w:pgNumType"))  # type: ignore[attr-defined]
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        section._sectPr.append(pg_num_type)  # type: ignore[attr-defined]
    pg_num_type.set(qn("w:start"), "1")


def add_cover(doc: DocumentType) -> None:
    for _ in range(3):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(8)
    run = kicker.add_run("OPEN ENGINEERING SPECIFICATION")
    set_run_font(run, size=10, bold=True, color=GOLD)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("ECK")
    set_run_font(run, size=40, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    run = subtitle.add_run("Embodied Cognitive Kernel")
    set_run_font(run, size=20, bold=True, color=TEAL)

    product = doc.add_paragraph()
    product.alignment = WD_ALIGN_PARAGRAPH.CENTER
    product.paragraph_format.space_after = Pt(18)
    run = product.add_run("Digital Life Kernel v0.1")
    set_run_font(run, size=15, color=BLUE)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.space_after = Pt(18)
    run = rule.add_run("CONTRACT-GROUNDED  •  VERIFIER-DRIVEN  •  LOCALLY PERSISTENT")
    set_run_font(run, size=9.5, bold=True, color=MUTED)
    set_bottom_border(rule, color=GOLD)

    statement = doc.add_paragraph()
    statement.alignment = WD_ALIGN_PARAGRAPH.CENTER
    statement.paragraph_format.left_indent = Inches(0.65)
    statement.paragraph_format.right_indent = Inches(0.65)
    statement.paragraph_format.space_after = Pt(20)
    add_inline(
        statement,
        "第一個公開可用版本的目標不是「會思考的 AI」，而是可連續運行、"
        "可恢復身分、可累積經驗且具有生命週期的數位生命核心。",
        size=13,
        color=INK,
    )

    status = doc.add_table(rows=1, cols=3)
    status.alignment = WD_TABLE_ALIGNMENT.CENTER
    prevent_row_split(status.rows[0])
    set_table_width(status, 7200)
    remove_table_borders(status)
    for index, (headline, detail, fill) in enumerate(
        [
            ("VERSION", "0.1.0", PALE_BLUE),
            ("STATUS", "Public alpha", PALE_TEAL),
            ("LICENSE", "Apache-2.0", PALE_GOLD),
        ]
    ):
        cell = status.rows[0].cells[index]
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=120, bottom=120, start=140, end=140)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(headline)
        set_run_font(run, size=8, bold=True, color=MUTED)
        run.add_break()
        run = paragraph.add_run(detail)
        set_run_font(run, size=11, bold=True, color=NAVY)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(42)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(0)
    run = meta.add_run("Repository: eck-digital-life-kernel\n")
    set_run_font(run, size=9.5, bold=True, color=NAVY)
    run = meta.add_run(f"Formal specification generated {RELEASE_DATE}")
    set_run_font(run, size=9, color=MUTED)


def add_callout(doc: DocumentType, title: str, body: str, fill: str = PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    prevent_row_split(table.rows[0])
    set_table_width(table)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    set_run_font(run, size=10, bold=True, color=TEAL)
    run.add_break()
    add_inline(paragraph, body, size=10.5, color=INK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_front_matter(doc: DocumentType, report: dict[str, object]) -> None:
    heading = doc.add_paragraph()
    heading.style = doc.styles["Heading 1"]
    heading.paragraph_format.space_before = Pt(0)
    add_inline(heading, "文件控制", size=16, color=BLUE)

    control_rows = [
        ["正式名稱", "ECK — Embodied Cognitive Kernel"],
        ["版本", "Digital Life Kernel v0.1.0"],
        ["儲存庫", "eck-digital-life-kernel"],
        ["授權", "Apache License 2.0"],
        ["主要文件語言", "繁體中文（程式碼與 API 使用英文）"],
        ["目標環境", "Windows 10/11 + WSL2 + Docker Desktop"],
        ["模型策略", "Ollama 預設；Mock 用於確定性測試；不硬編碼模型"],
        ["權重策略", "v0.1 不自動修改基礎模型權重"],
    ]
    add_table(doc, control_rows, header=False, widths=[2300, 7060])

    add_callout(
        doc,
        "誠實性邊界",
        "本文件嚴格區分 Implemented、Experimental、Future 與 Research。"
        "被跳過的環境檢查不算通過；GridWorld 經驗重用不宣稱抽象泛化；"
        "研究論文結果不得直接外推至 16GB VRAM 的目標電腦。",
        fill=PALE_GOLD,
    )

    heading = doc.add_paragraph()
    heading.style = doc.styles["Heading 2"]
    add_inline(heading, "v0.1 工程狀態矩陣", size=13, color=BLUE)
    matrix = [
        ["範圍", "狀態", "v0.1 證據"],
        ["生命週期與重啟恢復", "Implemented", "整合測試 + live HTTP 驗收"],
        ["Success Contract／Policy Gate", "Implemented", "單元測試 + 拒絕路徑"],
        [
            "Experience／Knowledge／Reflection／Skill",
            "Implemented",
            "證據准入、固定反思與整合測試",
        ],
        ["CLI／REST API／Web Dashboard", "Implemented", "live HTTP 驗收"],
        ["Ollama Provider", "Implemented", "介面與 Mocked HTTP 測試"],
        ["Rust event-chain verifier", "Experimental", "原始碼與測試；本機未執行"],
        ["Planner／Reflection／Curiosity", "Future / Research", "只有規格，未實作"],
        ["Prediction／通用世界模型", "Future / Research", "只有規格，未實作"],
        ["自動權重更新", "Prohibited", "v0.1 明確禁止"],
    ]
    add_table(doc, matrix, header=True, widths=[3150, 1700, 4510])

    heading = doc.add_paragraph()
    heading.style = doc.styles["Heading 2"]
    add_inline(heading, "目標開發電腦", size=13, color=BLUE)
    hardware = [
        ["元件", "規格"],
        ["CPU", "AMD Ryzen 7 9800X3D，8 核心／16 執行緒，Zen 5 3D V-Cache"],
        ["RAM", "64GB DDR5-6000 CL30（32GB × 2）"],
        ["GPU", "NVIDIA GeForce RTX 4070 Ti SUPER，16GB GDDR6X"],
        ["SSD", "Samsung 990 Pro／Kingston KC3000 2TB PCIe 4.0 NVMe"],
        ["ECK 空間", "約 1.2–1.5TB，保留 Windows 與其他軟體空間"],
        ["OS", "Windows 11；執行環境為 WSL2 + Docker Desktop"],
    ]
    add_table(doc, hardware, header=True, widths=[1900, 7460])

    heading = doc.add_paragraph()
    heading.style = doc.styles["Heading 1"]
    add_inline(heading, "驗收快照", size=16, color=BLUE)
    required_status = str(report.get("required_status", "unknown"))
    live = report.get("live_acceptance")
    live_status = live.get("status", "unknown") if isinstance(live, dict) else "unknown"
    acceptance = [
        ["項目", "結果", "說明"],
        ["Required checks", required_status, "Ruff、MyPy、測試、coverage"],
        ["Live HTTP acceptance", str(live_status), "Health、事件鏈、持久化、安全程式、GridWorld"],
    ]
    optional = report.get("optional_environment_checks")
    if isinstance(optional, dict):
        for name in ("docker", "rust"):
            result = optional.get(name)
            if isinstance(result, dict):
                acceptance.append(
                    [
                        f"Optional: {name}",
                        str(result.get("status", "unknown")),
                        str(result.get("reason", "")),
                    ]
                )
    add_table(doc, acceptance, header=True, widths=[2600, 1700, 5060])

    grid = live.get("gridworld") if isinstance(live, dict) else None
    if isinstance(grid, dict):
        add_callout(
            doc,
            "GridWorld 驗收限制",
            f"首次探索 {grid.get('first_exploration_steps')} steps，經驗重用後 "
            f"{grid.get('second_exploration_steps')} steps。"
            "這只證明同一環境 ID／幾何在表面標籤改變後的持久路徑重用，"
            "不構成抽象任務泛化證明。",
            fill=PALE_TEAL,
        )

    heading = doc.add_paragraph()
    heading.style = doc.styles["Heading 1"]
    add_inline(heading, "十三卷導覽", size=16, color=BLUE)
    overview = [["卷", "主題", "定位"]]
    for roman, title, filename in VOLUMES:
        status = extract_status(DOCS_DIR / filename)
        overview.append([roman, title, status])
    add_table(doc, overview, header=True, widths=[750, 4050, 4560])

    heading = doc.add_paragraph()
    heading.style = doc.styles["Heading 2"]
    add_inline(heading, "建議閱讀順序", size=13, color=BLUE)
    for text in [
        "決策者與協作者：卷 I → II → XII → XIII。",
        "核心開發者：卷 III → IV → V → IX → XI。",
        "部署與操作：README → 卷 VI → X → XII。",
        "研究協作者：卷 VII → VIII → XIII；先確認狀態標記與反證條件。",
    ]:
        paragraph = doc.add_paragraph(style="List Bullet")
        add_inline(paragraph, text)


def extract_status(path: Path) -> str:
    for line in path.read_text(encoding="utf-8").splitlines()[:10]:
        if line.startswith("**狀態：**"):
            return line.removeprefix("**狀態：**").strip()
    return "未標示"


def add_table(
    doc: DocumentType,
    rows: list[list[str]],
    *,
    header: bool,
    widths: list[int] | None = None,
) -> Table:
    if not rows:
        raise ValueError("table must contain at least one row")
    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)

    if widths is None:
        maxima = []
        for index in range(column_count):
            longest = max(len(row[index]) if index < len(row) else 0 for row in rows)
            maxima.append(max(10, min(longest, 42)))
        total = sum(maxima)
        widths = [int(USABLE_WIDTH_DXA * value / total) for value in maxima]
        widths[-1] += USABLE_WIDTH_DXA - sum(widths)

    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        if row_index == 0 and header:
            repeat_table_header(table.rows[-1])
        for index, cell in enumerate(cells):
            width = widths[index] if index < len(widths) else int(USABLE_WIDTH_DXA / column_count)
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0 and header:
                set_cell_shading(cell, PALE_BLUE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, "F7F9FB")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(1.5)
            paragraph.paragraph_format.line_spacing = 1.08
            text = values[index] if index < len(values) else ""
            add_inline(
                paragraph,
                text,
                size=8.5 if column_count >= 3 else 9,
                color=NAVY if row_index == 0 and header else INK,
            )
            if row_index == 0 and header:
                for run in paragraph.runs:
                    run.bold = True
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_rule(doc: DocumentType) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    set_bottom_border(paragraph, color=GRID)


def add_code_block(doc: DocumentType, lines: Iterable[str], language: str) -> None:
    paragraph = doc.add_paragraph(style="ECK Code")
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F4F6F8")
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), TEAL)
    borders.append(left)
    p_pr.append(borders)
    if language:
        run = paragraph.add_run(f"[{language}]\n")
        set_run_font(run, size=7.5, bold=True, color=TEAL, mono=True)
    content = "\n".join(lines)
    run = paragraph.add_run(content if content else " ")
    set_run_font(run, size=8, color=INK, mono=True)


def add_blockquote(doc: DocumentType, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    prevent_row_split(table.rows[0])
    set_table_width(table)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_GOLD)
    set_cell_margins(cell, top=110, bottom=110, start=180, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    add_inline(paragraph, text, size=10.5, color=INK)
    for run in paragraph.runs:
        run.italic = True


def clean_cell(value: str) -> str:
    value = value.strip()
    value = re.sub(r"^\||\|$", "", value).strip()
    value = re.sub(r"`([^`]+)`", r"\1", value)
    value = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
    value = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", value)
    return value


def parse_table(lines: list[str]) -> list[list[str]]:
    parsed = []
    for index, line in enumerate(lines):
        if index == 1 and re.fullmatch(r"\|?[\s:|-]+\|?", line.strip()):
            continue
        parsed.append([clean_cell(cell) for cell in line.strip().strip("|").split("|")])
    return parsed


def render_markdown(doc: DocumentType, path: Path, *, skip_title: bool = True) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    index = 0
    paragraph_buffer: list[str] = []
    active_ordered_num_id: int | None = None
    active_bullet_num_id: int | None = None

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer).strip()
        paragraph_buffer.clear()
        if text:
            paragraph = doc.add_paragraph()
            add_inline(paragraph, text)

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            active_ordered_num_id = None
            active_bullet_num_id = None
            index += 1
            continue
        if stripped.startswith("```"):
            flush_paragraph()
            active_ordered_num_id = None
            active_bullet_num_id = None
            language = stripped[3:].strip()
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(doc, code_lines, language)
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines):
            separator = lines[index + 1].strip()
            if separator.startswith("|") and re.fullmatch(r"\|?[\s:|-]+\|?", separator):
                flush_paragraph()
                active_ordered_num_id = None
                active_bullet_num_id = None
                table_lines = [line, lines[index + 1]]
                index += 2
                while index < len(lines) and lines[index].strip().startswith("|"):
                    table_lines.append(lines[index])
                    index += 1
                add_table(doc, parse_table(table_lines), header=True)
                continue
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            active_ordered_num_id = None
            active_bullet_num_id = None
            level = len(heading_match.group(1))
            if skip_title and level == 1:
                index += 1
                continue
            mapped_level = min(max(level - 1 if skip_title else level, 1), 4)
            paragraph = doc.add_paragraph(style=f"Heading {mapped_level}")
            add_inline(
                paragraph,
                heading_match.group(2),
                size={1: 16, 2: 13, 3: 12, 4: 11}[mapped_level],
                color=BLUE if mapped_level <= 2 else TEAL,
            )
            index += 1
            continue
        if stripped == "---":
            flush_paragraph()
            active_ordered_num_id = None
            active_bullet_num_id = None
            add_rule(doc)
            index += 1
            continue
        if stripped.startswith(">"):
            flush_paragraph()
            active_ordered_num_id = None
            active_bullet_num_id = None
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip()[1:].strip())
                index += 1
            add_blockquote(doc, " ".join(quote_lines))
            continue
        ordered_match = re.match(r"^(\s*)(\d+\.)(?:\s+|(?=[^\d\s]))(.+)$", line)
        if ordered_match:
            flush_paragraph()
            indent = min(len(ordered_match.group(1)) // 2, 2)
            if active_ordered_num_id is None:
                start = int(ordered_match.group(2)[:-1])
                active_ordered_num_id = create_ordered_list_numbering(doc, start=start)
            active_bullet_num_id = None
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.375 + 0.25 * indent)
            paragraph.paragraph_format.first_line_indent = Inches(-0.188)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.25
            apply_ordered_list_numbering(paragraph, active_ordered_num_id, indent)
            add_inline(paragraph, ordered_match.group(3))
            index += 1
            continue
        bullet_match = re.match(
            r"^(\s*)(?:-(?:\s+|(?=[^\s-]))|[*+]\s+)(.+)$",
            line,
        )
        if bullet_match:
            flush_paragraph()
            active_ordered_num_id = None
            if active_bullet_num_id is None:
                active_bullet_num_id = create_bullet_list_numbering(doc)
            indent = min(len(bullet_match.group(1)) // 2, 2)
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.375 + 0.25 * indent)
            paragraph.paragraph_format.first_line_indent = Inches(-0.188)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.25
            apply_ordered_list_numbering(paragraph, active_bullet_num_id, indent)
            add_inline(paragraph, bullet_match.group(2))
            index += 1
            continue
        active_ordered_num_id = None
        active_bullet_num_id = None
        paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()


def add_volume(doc: DocumentType, roman: str, title: str, path: Path) -> None:
    doc.add_page_break()
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(4)
    kicker.paragraph_format.space_after = Pt(5)
    run = kicker.add_run(f"VOLUME {roman}")
    set_run_font(run, size=9.5, bold=True, color=GOLD)

    heading = doc.add_paragraph()
    heading.paragraph_format.space_after = Pt(7)
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run(title)
    set_run_font(run, size=25, bold=True, color=NAVY)

    status = extract_status(path)
    status_line = doc.add_paragraph()
    status_line.paragraph_format.space_after = Pt(14)
    run = status_line.add_run("STATUS  ")
    set_run_font(run, size=8.5, bold=True, color=MUTED)
    run = status_line.add_run(status)
    set_run_font(run, size=10, bold=True, color=TEAL)
    set_bottom_border(status_line, color=GOLD)

    render_markdown(doc, path)


def add_adrs(doc: DocumentType) -> None:
    doc.add_page_break()
    heading = doc.add_paragraph()
    heading.paragraph_format.space_after = Pt(6)
    run = heading.add_run("Architecture Decision Records")
    set_run_font(run, size=25, bold=True, color=NAVY)
    intro = doc.add_paragraph()
    add_inline(
        intro,
        "以下 ADR 記錄 v0.1 已接受的重大架構決策。變更決策時應新增 ADR，"
        "不得只修改結論而抹除歷史脈絡。",
    )
    for filename in ADR_FILES:
        add_rule(doc)
        render_markdown(doc, DOCS_DIR / "adr" / filename, skip_title=False)


def add_release_appendix(doc: DocumentType, report: dict[str, object]) -> None:
    doc.add_page_break()
    heading = doc.add_paragraph()
    heading.paragraph_format.space_after = Pt(6)
    run = heading.add_run("Release Evidence Appendix")
    set_run_font(run, size=25, bold=True, color=NAVY)
    paragraph = doc.add_paragraph()
    add_inline(
        paragraph,
        "本附錄由 artifacts/release-report.json 產生。輸出只列出實際結果；"
        "skipped 不視為 passed。",
    )
    checks = report.get("checks")
    rows = [["檢查", "狀態", "實際摘要"]]
    if isinstance(checks, list):
        for check in checks:
            if not isinstance(check, dict):
                continue
            tail = str(check.get("output_tail", "")).strip().splitlines()
            summary = tail[-1] if tail else ""
            if check.get("name") == "tests":
                summary = next((line for line in tail if "[100%]" in line), summary)
            if check.get("name") == "coverage":
                summary = next((line for line in tail if line.startswith("TOTAL")), summary)
            rows.append(
                [
                    str(check.get("name", "")),
                    str(check.get("status", "")),
                    summary,
                ]
            )
    add_table(doc, rows, header=True, widths=[1900, 1500, 5960])

    report_meta = [
        ["欄位", "值"],
        ["schema_version", str(report.get("schema_version", ""))],
        ["version", str(report.get("version", ""))],
        ["generated_at", str(report.get("generated_at", ""))],
        ["required_status", str(report.get("required_status", ""))],
        ["source_tree_sha256", str(report.get("source_tree_sha256", ""))],
    ]
    add_table(doc, report_meta, header=True, widths=[2400, 6960])

    add_callout(
        doc,
        "Truthfulness note",
        str(report.get("truthfulness_note", "")),
        fill=PALE_GOLD,
    )


def remove_empty_first_paragraph(doc: DocumentType) -> None:
    if not doc.paragraphs:
        return
    paragraph = doc.paragraphs[0]
    if paragraph.text:
        return
    element = paragraph._element
    element.getparent().remove(element)


def build(output: Path) -> None:
    report = json.loads(RELEASE_REPORT.read_text(encoding="utf-8"))
    doc = Document()
    remove_empty_first_paragraph(doc)
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
    doc.core_properties.title = "ECK — Embodied Cognitive Kernel: Digital Life Kernel v0.1"
    doc.core_properties.subject = "Formal engineering specification and release evidence"
    doc.core_properties.author = "ECK Project Contributors"
    doc.core_properties.keywords = (
        "ECK, Digital Life Kernel, Success Contract, verifier-grounded learning"
    )
    doc.core_properties.comments = "Generated from the repository's 13-volume specification."

    add_cover(doc)
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_section)
    configure_header_footer(body_section)
    add_front_matter(doc, report)

    for roman, title, filename in VOLUMES:
        add_volume(doc, roman, title, DOCS_DIR / filename)
    add_adrs(doc)
    add_release_appendix(doc, report)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output",
        type=Path,
        default=PROJECT_ROOT.parent
        / "deliverables"
        / "ECK_Digital_Life_Kernel_v0.1_Project_Specification_zh-TW.docx",
    )
    args = parser.parse_args()
    build(args.output.resolve())
    print(args.output.resolve())


if __name__ == "__main__":
    main()
