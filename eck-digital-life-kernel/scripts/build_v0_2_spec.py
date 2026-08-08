"""Build the ECK v0.2 development specification from the retained v0.1 DOCX."""

from __future__ import annotations

import hashlib
import tempfile
from pathlib import Path

from build_project_doc import (
    BLUE,
    GOLD,
    INK,
    MUTED,
    NAVY,
    PALE_BLUE,
    PALE_GOLD,
    PALE_TEAL,
    TEAL,
    add_inline,
    add_page_field,
    add_table,
    configure_section,
    configure_styles,
    prevent_row_split,
    remove_table_borders,
    render_markdown,
    set_bottom_border,
    set_cell_margins,
    set_cell_shading,
    set_run_font,
    set_table_width,
)
from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

PROJECT_ROOT = Path(__file__).resolve().parents[1]
REFERENCE = PROJECT_ROOT.parents[1] / (
    "ECK_Digital_Life_Kernel_v0.1_Project_Specification_zh-TW.docx"
)
SOURCE = PROJECT_ROOT / "docs" / "15-v0.2.0-development-specification.md"
OUTPUT = PROJECT_ROOT / "docs" / (
    "ECK_Digital_Life_Kernel_v0.2.0_Development_Specification_zh-TW.docx"
)
REFERENCE_SHA256 = "0702B01F55198A6A6167D238F6AF856E984F56B4F49BE2F6FA0A3FFFC2F0ED67"
BUILD_DATE = "2026-08-08"


def clear_body(doc: DocumentType) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def clear_story(story: object) -> None:
    element = story._element  # type: ignore[attr-defined]
    for child in list(element):
        element.remove(child)


def configure_cover_section(section: object) -> None:
    configure_section(section)
    clear_story(section.header)  # type: ignore[attr-defined]
    clear_story(section.footer)  # type: ignore[attr-defined]


def configure_body_header_footer(section: object) -> None:
    configure_section(section)
    header = section.header  # type: ignore[attr-defined]
    header.is_linked_to_previous = False
    clear_story(header)
    paragraph = header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("ECK")
    set_run_font(run, size=9, bold=True, color=TEAL)
    run = paragraph.add_run("  |  Contract-Guided Autonomous Learning")
    set_run_font(run, size=8.5, color=MUTED)
    set_bottom_border(paragraph)

    footer = section.footer  # type: ignore[attr-defined]
    footer.is_linked_to_previous = False
    clear_story(footer)
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    remove_table_borders(table)
    left, right = table.rows[0].cells
    left.width = Inches(5.7)
    right.width = Inches(0.8)
    left_paragraph = left.paragraphs[0]
    left_paragraph.paragraph_format.space_after = Pt(0)
    run = left_paragraph.add_run(
        f"Digital Life Kernel v0.2 Development  •  Apache-2.0  •  {BUILD_DATE}"
    )
    set_run_font(run, size=8.5, color=MUTED)
    right_paragraph = right.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_paragraph.paragraph_format.space_after = Pt(0)
    add_page_field(right_paragraph)

    page_number = section._sectPr.find(qn("w:pgNumType"))  # type: ignore[attr-defined]
    if page_number is None:
        page_number = OxmlElement("w:pgNumType")
        section._sectPr.append(page_number)  # type: ignore[attr-defined]
    page_number.set(qn("w:start"), "1")


def add_cover(doc: DocumentType) -> None:
    for _ in range(3):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(8)
    run = kicker.add_run("DEVELOPMENT ENGINEERING SPECIFICATION")
    set_run_font(run, size=10, bold=True, color=GOLD)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("ECK")
    set_run_font(run, size=40, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    run = subtitle.add_run("Embodied Cognitive Kernel")
    set_run_font(run, size=20, bold=True, color=TEAL)

    product = doc.add_paragraph()
    product.alignment = WD_ALIGN_PARAGRAPH.CENTER
    product.paragraph_format.space_after = Pt(18)
    run = product.add_run("Digital Life Kernel v0.2.0")
    set_run_font(run, size=15, color=BLUE)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.space_after = Pt(18)
    run = rule.add_run(
        "CONTRACT-GUIDED  •  EVIDENCE-GROUNDED  •  CONTINUOUSLY EVALUATED"
    )
    set_run_font(run, size=9.5, bold=True, color=MUTED)
    set_bottom_border(rule, color=GOLD)

    statement = doc.add_paragraph()
    statement.alignment = WD_ALIGN_PARAGRAPH.CENTER
    statement.paragraph_format.left_indent = Inches(0.65)
    statement.paragraph_format.right_indent = Inches(0.65)
    statement.paragraph_format.space_after = Pt(20)
    add_inline(
        statement,
        "從固定能力走向可驗證的能力成長閉環：提出候選、取得證據、接受反例、"
        "產生技能、比較基線、通過保留集，並在失敗時安全回復。",
        size=13,
        color=INK,
    )

    status = doc.add_table(rows=1, cols=3)
    status.alignment = WD_TABLE_ALIGNMENT.CENTER
    prevent_row_split(status.rows[0])
    set_table_width(status, 7200)
    remove_table_borders(status)
    values = [
        ("VERSION", "0.2.0", PALE_BLUE),
        ("STATUS", "Development", PALE_TEAL),
        ("RUNTIME", "v0.1.0 verified", PALE_GOLD),
    ]
    for index, (headline, detail, fill) in enumerate(values):
        cell = status.rows[0].cells[index]
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=120, bottom=120, start=140, end=140)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(headline)
        set_run_font(run, size=8, bold=True, color=MUTED)
        run.add_break()
        run = paragraph.add_run(detail)
        set_run_font(run, size=11, bold=True, color=NAVY)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(42)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(0)
    run = meta.add_run("Repository: eck-digital-life-kernel\n")
    set_run_font(run, size=9.5, bold=True, color=NAVY)
    run = meta.add_run(f"Development specification generated {BUILD_DATE}")
    set_run_font(run, size=9, color=MUTED)


def add_front_matter(doc: DocumentType) -> None:
    heading = doc.add_paragraph(style="Heading 1")
    add_inline(heading, "文件控制", size=16, color=BLUE)
    add_table(
        doc,
        [
            ["欄位", "內容"],
            ["文件名稱", "ECK Digital Life Kernel v0.2.0 開發規格"],
            ["文件狀態", "Development Baseline；尚未完成 release gate"],
            ["前置版本", "v0.1.0 已通過並推送 GitHub main"],
            ["目前執行版本", "v0.1.0；不得提前宣稱 v0.2.0 已發布"],
            ["核心主題", "Contract-Guided Autonomous Learning"],
            ["硬體基線", "Windows 11、RTX 3060、Docker Desktop、本機 Ollama"],
            ["授權", "Apache-2.0；外部模型與套件依各自條款"],
        ],
        header=True,
        widths=[2400, 6960],
    )

    heading = doc.add_paragraph(style="Heading 2")
    add_inline(heading, "版本真實性", size=13, color=BLUE)
    truth = doc.add_paragraph()
    add_inline(
        truth,
        "本文件代表 v0.2.0 已進入設計與開發，不代表功能已完成。套件、API 與"
        "執行中核心繼續維持 v0.1.0，直到本文件第 17 節的全部 release gate 通過。",
    )

    heading = doc.add_paragraph(style="Heading 2")
    add_inline(heading, "工程結果摘要", size=13, color=BLUE)
    add_table(
        doc,
        [
            ["階段", "主要交付", "證明方式"],
            ["M2.0", "最新資訊批判學習", "來源快照、主張證據、無結論品質閘門"],
            ["M2.1", "Work Portfolio 與學習產率", "停滯與重複工作故障注入"],
            ["M2.2", "Contract-Guided Planner", "固定合約、bounded repair、基線比較"],
            ["M2.3", "Experience Graph", "來源、反例、技能與結果可追溯"],
            ["M2.4", "Capability Gap Forge", "三類技能隔離測試、熱啟用、回復"],
            ["M2.5", "Evaluation 與 Model Lab", "保留集、能力曲線、人工權重升級閘門"],
        ],
        header=True,
        widths=[1200, 3380, 4780],
    )


def enable_field_updates(doc: DocumentType) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main() -> None:
    if not REFERENCE.is_file():
        raise FileNotFoundError(f"Reference DOCX not found: {REFERENCE}")
    digest = hashlib.sha256(REFERENCE.read_bytes()).hexdigest().upper()
    if digest != REFERENCE_SHA256:
        raise RuntimeError("The retained v0.1 reference DOCX hash changed.")

    doc = Document(REFERENCE)
    clear_body(doc)
    configure_styles(doc)
    configure_cover_section(doc.sections[0])
    add_cover(doc)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_body_header_footer(body_section)
    add_front_matter(doc)
    source_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    first_section = next(
        index for index, line in enumerate(source_lines) if line.startswith("## 1.")
    )
    with tempfile.TemporaryDirectory(prefix="eck-v02-doc-") as temp_dir:
        body_source = Path(temp_dir) / "body.md"
        body_source.write_text(
            "\n".join(source_lines[first_section:]) + "\n",
            encoding="utf-8",
        )
        render_markdown(doc, body_source, skip_title=True)
    enable_field_updates(doc)

    doc.core_properties.title = "ECK Digital Life Kernel v0.2.0 Development Specification"
    doc.core_properties.subject = "Contract-Guided Autonomous Learning"
    doc.core_properties.author = "ECK Project Contributors"
    doc.core_properties.keywords = (
        "ECK, autonomous learning, success contract, experience graph, skill forge"
    )
    doc.core_properties.comments = (
        "Template-derived from the retained ECK v0.1 engineering specification."
    )
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
